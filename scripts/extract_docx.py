#!/usr/bin/env python3
"""
从 Word (.docx) 文件中提取文字内容和格式信息。
用法: python extract_docx.py <input.docx> [output.json]
"""
import sys
import json
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("请安装 python-docx: pip install python-docx")
    sys.exit(1)


def get_run_formatting(run) -> dict:
    """获取文本格式信息"""
    fmt = {}
    if run.bold:
        fmt["bold"] = True
    if run.italic:
        fmt["italic"] = True
    if run.underline:
        fmt["underline"] = True
    if run.font.color and run.font.color.rgb:
        fmt["color"] = str(run.font.color.rgb)
    if run.font.highlight_color is not None:
        fmt["highlight"] = str(run.font.highlight_color)
    if run.font.size:
        fmt["size"] = run.font.size.pt
    return fmt


def is_key_point(fmt: dict) -> bool:
    """判断是否为老师标注的重点"""
    color = fmt.get("color", "")
    highlight = fmt.get("highlight", "")

    # 红色/黄色字体或高亮 → 重点
    red_colors = ["FF0000", "FF000000", "C00000", "E60000"]
    yellow_highlights = ["YELLOW", "7", "yellow"]

    if color and any(c in str(color).upper() for c in red_colors):
        return True
    if highlight and any(h in str(highlight).upper() for h in yellow_highlights):
        return True
    return False


def extract_docx(docx_path: str) -> dict:
    doc = Document(docx_path)
    elements = []
    key_points = []

    for para in doc.paragraphs:
        para_data = {"type": "paragraph", "style": para.style.name if para.style else "", "runs": []}
        full_text = ""

        for run in para.runs:
            text = run.text
            if not text.strip():
                continue
            fmt = get_run_formatting(run)
            full_text += text
            run_data = {"text": text, "formatting": fmt}
            if is_key_point(fmt):
                run_data["is_key_point"] = True
                key_points.append(f"★ [段落: {para.style.name}] {full_text.strip()}")
            para_data["runs"].append(run_data)

        para_data["full_text"] = full_text.strip()
        if full_text.strip():
            elements.append(para_data)

    # 提取表格内容
    tables_data = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables_data.append(rows)

    return {
        "total_paragraphs": len([e for e in elements if e["full_text"]]),
        "total_tables": len(tables_data),
        "content": elements,
        "tables": tables_data,
        "key_points": key_points,
        "key_point_count": len(key_points),
        "full_text": "\n\n".join(e["full_text"] for e in elements if e["full_text"]),
        "full_text_with_formatting": "\n\n".join(
            f"[{e['style']}] {e['full_text']}" for e in elements if e["full_text"]
        )
    }


def main():
    if len(sys.argv) < 2:
        print("用法: python extract_docx.py <input.docx> [output.json]")
        sys.exit(1)

    docx_path = sys.argv[1]
    if not Path(docx_path).exists():
        print(f"错误: 文件不存在 — {docx_path}")
        sys.exit(1)

    result = extract_docx(docx_path)

    if len(sys.argv) > 2:
        with open(sys.argv[2], "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"结果已保存到: {sys.argv[2]}")

    print(f"\n=== 文档统计 ===")
    print(f"段落数: {result['total_paragraphs']}")
    print(f"表格数: {result['total_tables']}")
    print(f"检测到重点标记: {result['key_point_count']} 处")

    if result["key_points"]:
        print(f"\n=== 重点内容 ===")
        for kp in result["key_points"][:10]:
            print(kp)
        if len(result["key_points"]) > 10:
            print(f"... 共 {len(result['key_points'])} 处")

    print(f"\n=== 文档全文 ===")
    print(result["full_text"][:3000])
    if len(result["full_text"]) > 3000:
        print(f"\n... (共 {len(result['full_text'])} 字，仅显示前 3000)")


if __name__ == "__main__":
    main()
